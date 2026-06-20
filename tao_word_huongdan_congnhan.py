import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Helper: set cell shading
def set_cell_shading(cell, color):
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shd)

# Helper: add formatted paragraph
def add_para(text, bold=False, italic=False, size=11, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6, bullet=False):
    style = 'List Bullet' if bullet else 'Normal'
    p = doc.add_paragraph(style=style)
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Segoe UI'
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

# Helper: add heading
def add_heading_styled(text, level, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Segoe UI'
    
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(11, 93, 59) # xanh dam
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(18, 122, 78) # xanh nhat
    else:
        run.font.size = Pt(11.5)
        run.font.color.rgb = RGBColor(91, 33, 182) # tim
    return p

# Helper: add alert box (callout)
def add_callout(text, title="Lưu ý:"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.rows[0].cells[0]
    cell.width = Cm(16)
    set_cell_shading(cell, 'E7F3ED') # xanh nhat
    
    # Left border styling in XML
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(r'<w:tcBorders %s><w:left w:val="single" w:sz="24" w:space="0" w:color="0B5D3B"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>' % nsdecls('w'))
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    
    run_title = p.add_run(f"{title} ")
    run_title.bold = True
    run_title.font.name = 'Segoe UI'
    run_title.font.size = Pt(10.5)
    run_title.font.color.rgb = RGBColor(11, 93, 59)
    
    run_text = p.add_run(text)
    run_text.italic = True
    run_text.font.name = 'Segoe UI'
    run_text.font.size = Pt(10)
    run_text.font.color.rgb = RGBColor(31, 41, 55)

# ========== TRANG BÌA SANG TRỌNG ==========
doc.add_paragraph()
doc.add_paragraph()
add_para('CÔNG TY TNHH TM-SX HUỆ LINH', bold=True, size=13, color=(11,93,59), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para('—' * 30, size=10, color=(180,180,180), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
doc.add_paragraph()
doc.add_paragraph()
add_para('HƯỚNG DẪN THAO TÁC NHANH', bold=True, size=24, color=(11,93,59), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_para('& XEM BÁO CÁO KPI V2.0', bold=True, size=20, color=(11,93,59), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
doc.add_paragraph()
add_para('Dành cho Tổ trưởng, Quản lý và Công nhân sản xuất', italic=True, size=12, color=(107,114,128), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
doc.add_paragraph()
doc.add_paragraph()
add_para('Mã tài liệu: 001-QLHS001-CN', size=11, color=(107,114,128), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para('Phiên bản: V2.0', size=11, color=(107,114,128), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para('Tháng 06 năm 2026', size=11, color=(107,114,128), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

doc.add_page_break()

# ========== NỘI DUNG CHÍNH ==========
add_heading_styled('HƯỚNG DẪN THAO TÁC NHANH & XEM BÁO CÁO', 1)
add_para('Tài liệu này hướng dẫn anh chị em cách đăng nhập, thao tác và lọc xem các báo cáo kết quả làm việc của mình trên hệ thống một cách dễ dàng và nhanh chóng nhất.', italic=True, size=11, color=(75,85,99), space_after=12)

# --- PHẦN 1 ---
add_heading_styled('🔑 PHẦN 1: ĐĂNG NHẬP VÀO HỆ THỐNG', 2)
add_para('1. Địa chỉ truy cập: Mở ứng dụng trên trình duyệt web của máy tính hoặc điện thoại.', space_after=4)
add_para('2. Tên đăng nhập (Username):', space_after=3)
add_para('Tổ trưởng / Quản lý: Nhập tài khoản quản lý được cấp (Ví dụ: admin, quanly).', bullet=True, space_after=2)
add_para('Công nhân: Nhập chính xác Mã nhân viên của mình (Ví dụ: T0006, T0024,...).', bullet=True, space_after=3)
add_para('3. Mật khẩu: Nhập mật khẩu của anh chị (mặc định ban đầu của admin là admin123).', space_after=4)
add_para('4. Bấm nút Đăng nhập. Khi có thông báo “Chào mừng...” màu xanh xuất hiện là thành công.', space_after=8)

# --- PHẦN 2 ---
add_heading_styled('📝 PHẦN 2: CÁCH ĐỌC & HIỂU ĐIỂM SỐ CỦA BẢN THÂN', 2)
add_para('Hệ thống đánh giá kết quả làm việc của anh chị em dựa trên thang điểm từ 1 đến 5:', space_after=6)
add_para('🔴 Từ 1.0 đến dưới 2.0 (Không đạt): Làm việc chưa đúng quy trình, cần được đào tạo lại công đoạn đó.', bullet=True, space_after=3)
add_para('🟡 Từ 2.0 đến dưới 3.0 (Cần cải thiện): Đã biết làm nhưng thao tác còn chậm, hay quên vệ sinh máy hoặc đi muộn.', bullet=True, space_after=3)
add_para('🔵 Từ 3.0 đến dưới 3.8 (Đạt yêu cầu): Làm tốt, đúng giờ, sạch sẽ, đạt năng suất ca được giao. (Mốc chuẩn)', bullet=True, space_after=3)
add_para('🟢 Từ 3.8 đến dưới 4.5 (Tốt): Làm việc xuất sắc, chủ động giúp đỡ tổ viên khác khi rảnh.', bullet=True, space_after=3)
add_para('❇️ Từ 4.5 đến 5.0 (Xuất sắc): Đạt sản lượng cao vượt trội, phát hiện kịp thời lỗi máy móc hoặc đề xuất cải tiến có lợi cho xưởng.', bullet=True, space_after=8)

# --- PHẦN 3 ---
add_heading_styled('📈 PHẦN 3: CÁCH DÙNG BỘ LỌC ĐỂ XEM BÁO CÁO CHI TIẾT', 2)
add_para('Để xem điểm số của mình hoặc của tổ mình quản lý, anh chị em làm theo các bước sau:', space_after=6)

add_heading_styled('Bước 1: Vào mục Báo cáo nâng cao', 3)
add_para('Nhìn sang thanh trình đơn bên trái (Sidebar), bấm chọn biểu tượng 📈 Báo cáo nâng cao.', space_after=4)

add_heading_styled('Bước 2: Chọn loại phiếu muốn xem', 3)
add_para('Ở trên cùng màn hình báo cáo sẽ có 2 nút lớn:', space_after=3)
add_para('Đánh giá năng lực: Xem tổng kết năng lực tay nghề định kỳ (Đa kỹ năng, vận hành máy,...).', bullet=True, space_after=2)
add_para('Đánh giá hàng ngày: Xem điểm chấm công việc, 5S và kỷ luật hàng ngày của từng ca.', bullet=True, space_after=4)

add_heading_styled('Bước 3: Sử dụng các ô lọc (Rất quan trọng!)', 3)
add_para('Anh chị có thể kết hợp các ô lọc để tìm đúng thông tin cần thiết:', space_after=4)
add_para('Từ ngày / Đến ngày: Chọn khoảng thời gian muốn xem (Ví dụ: từ ngày 01-06-2026 đến 15-06-2026).', bullet=True, space_after=2)
add_para('Máy: Chọn đúng máy phụ trách (Ví dụ: Cán 01, In 02,...).', bullet=True, space_after=2)
add_para('Công đoạn: Lọc theo khâu sản xuất (Ví dụ: Cán, In, Ghép,...).', bullet=True, space_after=2)
add_para('Phòng ban: Chọn bộ phận làm việc (Ví dụ: MÁY CÁN, IN, DỆT,...).', bullet=True, space_after=2)
add_para('Nhân viên: Bấm vào đây và chọn đúng tên của mình để xem riêng điểm của bản thân.', bullet=True, space_after=6)

add_callout('Nếu muốn xem lại toàn bộ dữ liệu từ đầu, hãy bấm nút Xóa lọc màu xám.', title='Mẹo:')

doc.add_page_break()

# --- PHẦN 4 ---
add_heading_styled('📊 PHẦN 4: CÁC THÔNG TIN TRÊN MÀN HÌNH BÁO CÁO', 2)
add_para('Sau khi chọn bộ lọc, màn hình sẽ hiển thị 4 khu vực chính:', space_after=6)
add_para('1. Hộp chỉ số tổng hợp (Màu xanh/cam):', bold=True, space_after=3)
add_para('Hiển thị tổng số ngày được chấm điểm.', bullet=True, space_after=2)
add_para('Điểm trung bình đạt được (kèm xếp loại đạt hay tốt).', bullet=True, space_after=2)
add_para('Tỷ lệ % số ngày đạt chuẩn.', bullet=True, space_after=4)
add_para('2. Biểu đồ xu hướng (Đường kẻ màu xanh):', bold=True, space_after=3)
add_para('Giúp anh chị nhìn thấy điểm số của mình đang tăng lên hay giảm đi theo từng ngày/tháng.', bullet=True, space_after=2)
add_para('Các chấm tròn trên biểu đồ hiển thị điểm số trung bình thực tế của mốc thời gian đó.', bullet=True, space_after=4)
add_para('3. Bảng danh sách xuất sắc & cần cải thiện:', bold=True, space_after=3)
add_para('Hiển thị danh sách những anh chị em đạt điểm cao nhất và những người đang cần hỗ trợ đào tạo thêm.', bullet=True, space_after=4)
add_para('4. Bảng chi tiết phía dưới:', bold=True, space_after=3)
add_para('Liệt kê chi tiết từng ngày đánh giá.', bullet=True, space_after=2)
add_para('Anh chị có thể bấm nút Xem ở cuối dòng để mở xem lại toàn bộ phiếu chấm điểm chi tiết của ngày hôm đó (có đầy đủ nhận xét của Quản đốc).', bullet=True, space_after=8)

# --- PHẦN 5 ---
add_heading_styled('🖨 PHẦN 5: XUẤT FILE VÀ IN BÁO CÁO', 2)
add_para('• In báo cáo (PDF): Bấm nút In báo cáo màu xanh lá. Giao diện in ấn của trình duyệt sẽ hiện ra, hệ thống đã tự động ẩn đi các nút bấm thừa để bản in ra giấy hoặc lưu thành file PDF được gọn gàng và đẹp nhất.', space_after=4)
add_para('• Tải file Excel (CSV): Bấm nút Xuất CSV chi tiết màu xám để tải dữ liệu về máy tính dưới dạng file Excel. File này đã được xử lý để khi mở bằng Excel trên máy tính sẽ hiển thị đúng chữ tiếng Việt có dấu, không bị lỗi font chữ.', space_after=12)

doc.add_paragraph()
add_para('Chúc anh chị em tổ sản xuất đạt kết quả thật cao và vận hành thiết bị an toàn!', bold=True, size=11, color=(11,93,59), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para('Nếu có thắc mắc trong quá trình thao tác, vui lòng liên hệ Tổ trưởng hoặc bộ phận Kỹ thuật để được hỗ trợ.', italic=True, size=10, color=(107,114,128), align=WD_ALIGN_PARAGRAPH.CENTER)

# Save the document
output_path = r'D:\Cty Hue Linh\Danh Muc\Quy_Trinh_Danh_Gia_HL\App_Danh_Gia_KPI_V2\Huong_Dan_Su_Dung_App_KPI.docx'
doc.save(output_path)
print(f"SUCCESS: File saved at {output_path}")
