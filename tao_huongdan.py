from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Helper: set cell shading
def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

# Helper: add formatted paragraph
def add_para(text, bold=False, size=11, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Segoe UI'
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

# Helper: styled table
def make_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Segoe UI'
        run.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '0B5D3B')
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'Segoe UI'
            if ri % 2 == 1:
                set_cell_shading(cell, 'E7F3ED')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

# ========== TRANG BIA ==========
doc.add_paragraph()
doc.add_paragraph()
add_para('CONG TY TNHH TM-SX HUE LINH', bold=True, size=14, color=(11,93,59), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para('_' * 50, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
doc.add_paragraph()
add_para('HUONG DAN SU DUNG', bold=True, size=24, color=(11,93,59), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
add_para('APP DANH GIA NANG LUC', bold=True, size=20, color=(11,93,59), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
add_para('CONG NHAN SAN XUAT', bold=True, size=20, color=(11,93,59), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
doc.add_paragraph()
add_para('Phien ban: V2.0', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para('Ma tai lieu: 001-QLHS001', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para('Ngay ban hanh: 17/06/2026', size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
doc.add_paragraph()
add_para('Ho tro: Android / iOS / Windows PC', size=11, color=(100,100,100), align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('Cong nghe: PWA (Progressive Web App) — Chay tren trinh duyet', size=11, color=(100,100,100), align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ========== MUC LUC ==========
add_para('MUC LUC', bold=True, size=16, color=(11,93,59), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)
muc_luc = [
    '1. Gioi thieu ung dung',
    '2. Yeu cau he thong',
    '3. Cach cai dat & truy cap',
    '   3.1. Tren dien thoai Android',
    '   3.2. Tren dien thoai iOS (iPhone/iPad)',
    '   3.3. Tren may tinh PC (Windows)',
    '4. Huong dan su dung tung chuc nang',
    '   4.1. Bang dieu khien (Dashboard)',
    '   4.2. Phieu danh gia nang luc',
    '   4.3. Lich su & Bao cao',
    '   4.4. Nhan su & Danh muc',
    '   4.5. Phieu ban giao ca',
    '   4.6. KPI Nha may',
    '   4.7. Bao tri may',
    '   4.8. Sao luu & Cai dat',
    '5. Cach deploy (dua len mang) de dung chung',
    '6. Cau hoi thuong gap (FAQ)',
    '7. Luu y quan trong',
]
for item in muc_luc:
    add_para(item, size=11, space_after=4)

doc.add_page_break()

# ========== 1. GIOI THIEU ==========
add_para('1. GIOI THIEU UNG DUNG', bold=True, size=14, color=(11,93,59), space_after=10)
add_para('App Danh gia Nang luc Cong nhan San xuat la ung dung web (PWA) duoc thiet ke rieng cho Cong ty TNHH TM-SX Hue Linh, giup:', size=11, space_after=8)
benefits = [
    'Danh gia nang luc cong nhan hang ngay theo 3 nhom tieu chi (A/B/C)',
    'Ghi nhan ban giao ca: ve sinh, san luong, phe pham, tinh trang may',
    'Theo doi KPI nha may: tong san luong, ty le phe pham, OEE',
    'Quan ly bao tri may: trang thai, lich dinh ky, lich su sua chua',
    'Xuat bao cao CSV/JSON, dong bo Google Sheets',
    'Hoat dong OFFLINE — khong can Internet khi da cai dat',
]
for b in benefits:
    add_para(f'  •  {b}', size=11, space_after=3)

doc.add_paragraph()

# ========== 2. YEU CAU HE THONG ==========
add_para('2. YEU CAU HE THONG', bold=True, size=14, color=(11,93,59), space_after=10)
make_table(
    ['Thiet bi', 'Trinh duyet ho tro', 'Ghi chu'],
    [
        ['Android', 'Chrome 80+, Edge, Samsung Internet', 'Khuyen dung Chrome'],
        ['iOS (iPhone/iPad)', 'Safari 14+', 'Bat buoc dung Safari de cai PWA'],
        ['Windows PC', 'Chrome, Edge, Firefox', 'Khuyen dung Chrome hoac Edge'],
        ['Mac', 'Chrome, Safari, Edge', 'Tuong tu Windows'],
    ],
    col_widths=[4, 6, 5]
)

doc.add_paragraph()

# ========== 3. CACH CAI DAT ==========
add_para('3. CACH CAI DAT & TRUY CAP', bold=True, size=14, color=(11,93,59), space_after=10)

add_para('3.1. Tren dien thoai Android', bold=True, size=12, color=(18,122,78), space_after=8)
steps_android = [
    'Mo trinh duyet Chrome',
    'Nhap dia chi app vao thanh URL (VD: https://ten-app.github.io hoac http://192.168.x.x:8080)',
    'Nhan nut 3 cham doc (goc tren phai)',
    'Chon "Them vao Man hinh chinh" (Add to Home screen)',
    'Dat ten: "KPI Hue Linh" -> nhan Them',
    'Icon KPI xuat hien tren man hinh chinh — mo len nhu app that!',
]
for i, s in enumerate(steps_android, 1):
    add_para(f'  Buoc {i}: {s}', size=11, space_after=3)

doc.add_paragraph()
add_para('3.2. Tren dien thoai iOS (iPhone / iPad)', bold=True, size=12, color=(18,122,78), space_after=8)
steps_ios = [
    'Mo trinh duyet Safari (bat buoc Safari, Chrome tren iOS khong ho tro PWA)',
    'Nhap dia chi app vao thanh URL',
    'Nhan nut Chia se (hinh vuong co mui ten len, thanh duoi cung)',
    'Cuon xuong, chon "Them vao MH chinh" (Add to Home Screen)',
    'Dat ten: "KPI Hue Linh" -> nhan Them',
    'Icon KPI xuat hien tren man hinh chinh — mo len toan man hinh!',
]
for i, s in enumerate(steps_ios, 1):
    add_para(f'  Buoc {i}: {s}', size=11, space_after=3)

doc.add_paragraph()
add_para('3.3. Tren may tinh PC (Windows)', bold=True, size=12, color=(18,122,78), space_after=8)
steps_pc = [
    'Mo Chrome hoac Microsoft Edge',
    'Nhap dia chi app (hoac mo truc tiep file index.html)',
    'Tren thanh dia chi se thay bieu tuong (+) "Cai dat ung dung"',
    'Hoac nhan nut 3 cham -> "Cai dat ung dung..." / "Install app"',
    'Nhan Cai dat — app mo ra cua so rieng, co shortcut tren Desktop',
]
for i, s in enumerate(steps_pc, 1):
    add_para(f'  Buoc {i}: {s}', size=11, space_after=3)

add_para('  * Cach nhanh (offline): Nhap doi vao file index.html trong thu muc App_Danh_Gia_KPI/', size=10, color=(100,100,100), space_after=6)

doc.add_page_break()

# ========== 4. HUONG DAN TUNG CHUC NANG ==========
add_para('4. HUONG DAN SU DUNG TUNG CHUC NANG', bold=True, size=14, color=(11,93,59), space_after=10)

# 4.1 Dashboard
add_para('4.1. Bang dieu khien (Dashboard)', bold=True, size=12, color=(18,122,78), space_after=8)
add_para('Day la man hinh chinh hien thi ngay khi mo app. Bao gom:', size=11, space_after=6)
dash_items = [
    'Diem nen trung binh (Nhom A) — toan bo phieu da luu',
    'So phieu / so nhan vien da danh gia',
    'Ty le dat chuan (diem >= 3.0)',
    'Da ky nang trung binh (nhom B+C)',
    'Bieu do diem theo tieu chi, theo may, theo phong ban',
    'Phan bo xep loai: Xuat sac / Tot / Dat / Can cai thien / Khong dat',
]
for item in dash_items:
    add_para(f'  •  {item}', size=11, space_after=2)

doc.add_paragraph()

# 4.2 Phieu danh gia
add_para('4.2. Phieu danh gia nang luc', bold=True, size=12, color=(18,122,78), space_after=8)
add_para('Quy trinh nhap phieu:', size=11, space_after=6)
eval_steps = [
    'Chon nhan vien tu dropdown (tim theo ten, ma, phong ban)',
    'Chon may, cong doan, ca, ngay ra soat, nguoi danh gia',
    'Cham diem NHOM A (5 tieu chi, thang 1-5): Ve sinh, Van hanh, Tuan thu, Ky luat, An toan',
    'Cham diem NHOM B (3 chi so, chi 4 hoac 5): Bao cao bat thuong, De xuat cai tien, Ngan ngua loi',
    'Cham diem NHOM C (3 chi so, chi 4 hoac 5): Van hanh nhieu may, Ho tro khau khac, Huong dan CN moi',
    'App TU DONG tinh diem nen A va de xuat ket qua tong the',
    'Chon Hanh dong nhan su & Anh huong cap bac luong',
    'Ghi y kien quan doc (neu can)',
    'Nhan "Luu phieu" — xong!',
]
for i, s in enumerate(eval_steps, 1):
    add_para(f'  {i}. {s}', size=11, space_after=3)

doc.add_paragraph()

# 4.3 Lich su
add_para('4.3. Lich su & Bao cao', bold=True, size=12, color=(18,122,78), space_after=8)
add_para('  •  Loc phieu theo: ngay, may, phong ban, ten nhan vien', size=11, space_after=3)
add_para('  •  Xem chi tiet phieu (in duoc)', size=11, space_after=3)
add_para('  •  Sua / Xoa phieu', size=11, space_after=3)
add_para('  •  In danh sach hoac in tung phieu rieng le', size=11, space_after=3)

doc.add_paragraph()

# 4.4 Nhan su
add_para('4.4. Nhan su & Danh muc', bold=True, size=12, color=(18,122,78), space_after=8)
add_para('Quan ly:', size=11, space_after=6)
add_para('  •  Danh sach nhan su (56 nguoi tu Excel, co the them/sua/xoa)', size=11, space_after=3)
add_para('  •  Danh muc may moc (Can 01, In 01, Ghep 02...)', size=11, space_after=3)
add_para('  •  Danh muc cong doan (Can, In, Ghep, Thanh pham)', size=11, space_after=3)
add_para('  •  Danh muc ca lam viec', size=11, space_after=3)
add_para('  •  Tuy chon ket qua / hanh dong nhan su / anh huong luong', size=11, space_after=3)

doc.add_paragraph()

# 4.5 Ban giao ca
add_para('4.5. Phieu ban giao ca (MOI)', bold=True, size=12, color=(18,122,78), space_after=8)
add_para('Ghi nhan tinh trang khi chuyen ca:', size=11, space_after=6)
make_table(
    ['Truong du lieu', 'Mo ta', 'Lua chon'],
    [
        ['Ngay', 'Ngay ban giao', 'Chon tu lich'],
        ['Ca giao / Ca nhan', 'Ca hien tai va ca ke tiep', 'Dropdown'],
        ['May', 'May dang phu trach', 'Dropdown tu danh muc'],
        ['Nguoi giao / Nguoi nhan', 'Ai giao, ai nhan', 'Dropdown nhan su'],
        ['Ve sinh may', 'Tinh trang ve sinh', 'Dat / Chua dat'],
        ['San luong', 'So cuon/tam san xuat trong ca', 'Nhap so'],
        ['Phe pham', 'So cuon/tam phe pham phat sinh', 'Nhap so'],
        ['Tinh trang may', 'May co van de gi khong', 'Tot / Loi nhe / Hong'],
        ['Ghi chu', 'Van de can luu y cho ca sau', 'Nhap text'],
    ],
    col_widths=[4, 5.5, 4.5]
)

doc.add_paragraph()

# 4.6 KPI Nha may
add_para('4.6. KPI Nha may (MOI)', bold=True, size=12, color=(18,122,78), space_after=8)
add_para('Tu dong tong hop tu Phieu ban giao ca + Lich su sua chua:', size=11, space_after=6)
make_table(
    ['Chi so KPI', 'Cong thuc', 'Muc tieu'],
    [
        ['Tong san luong', 'Cong tat ca san luong cac phieu BGC', 'Cang cao cang tot'],
        ['Ty le phe pham (%)', 'Tong PP / Tong SL x 100', '< 5% la tot'],
        ['Thoi gian dung may', 'Tong gio dung tu lich su sua chua', 'Cang thap cang tot'],
        ['OEE (Hieu suat TB)', '(Gio kha dung - Gio dung) / Gio kha dung x 100', '>= 80% la tot'],
    ],
    col_widths=[4.5, 6, 4]
)
add_para('  * Bieu do san luong 7 ngay gan nhat va bang san luong theo may cung hien thi.', size=10, color=(100,100,100), space_after=6)

doc.add_paragraph()

# 4.7 Bao tri may
add_para('4.7. Bao tri may (MOI)', bold=True, size=12, color=(18,122,78), space_after=8)
add_para('Gom 3 phan:', size=11, space_after=6)
add_para('  a) Trang thai may moc:', bold=True, size=11, space_after=4)
add_para('     Moi may co trang thai: Dang chay (xanh) / Dang sua (vang) / Dung-Hong (do)', size=11, space_after=3)
add_para('     Cap nhat truc tiep bang cach chon dropdown va nhan "Luu"', size=11, space_after=6)
add_para('  b) Lich bao tri dinh ky:', bold=True, size=11, space_after=4)
add_para('     Them muc bao tri: may, noi dung, chu ky (hang tuan/thang/quy/nam), lan gan nhat, lan toi', size=11, space_after=6)
add_para('  c) Lich su sua chua:', bold=True, size=11, space_after=4)
add_para('     Ghi nhan: ngay, may, mo ta su co, thoi gian dung may (gio), nguoi sua', size=11, space_after=3)
add_para('     Du lieu nay cung cap nguon cho KPI "Thoi gian dung may" va OEE', size=11, space_after=6)

doc.add_paragraph()

# 4.8 Sao luu
add_para('4.8. Sao luu & Cai dat', bold=True, size=12, color=(18,122,78), space_after=8)
add_para('  •  Che do Cuc bo (TH1): Luu tren thiet bi, offline 100%', size=11, space_after=3)
add_para('  •  Che do Google Sheets (TH2): Moi phieu gui len Sheet dung chung qua Apps Script', size=11, space_after=3)
add_para('  •  Xuat JSON: Sao luu toan bo du lieu (khoi phuc duoc)', size=11, space_after=3)
add_para('  •  Xuat CSV: Bang tong hop cho Excel (chi doc, khong khoi phuc)', size=11, space_after=3)
add_para('  •  Nhap JSON: Khoi phuc du lieu tu file sao luu', size=11, space_after=3)

doc.add_page_break()

# ========== 5. DEPLOY ==========
add_para('5. CACH DEPLOY (DUA LEN MANG) DE DUNG CHUNG', bold=True, size=14, color=(11,93,59), space_after=10)
add_para('Hien tai app chay LOCAL tren 1 thiet bi. De nhieu nguoi truy cap qua dien thoai, can host len mang:', size=11, space_after=8)

make_table(
    ['Phuong phap', 'Chi phi', 'Do kho', 'Ket qua'],
    [
        ['GitHub Pages', 'Mien phi', 'De', 'URL: https://xxx.github.io/kpi-hue-linh'],
        ['Netlify', 'Mien phi', 'Rat de (keo tha)', 'URL: https://kpi-huelinh.netlify.app'],
        ['Server noi bo', 'Co san', 'Trung binh', 'URL: http://192.168.x.x:8080'],
        ['Google Drive', 'Mien phi', 'De', 'Tai ve dung local'],
    ],
    col_widths=[3.5, 2.5, 3, 5.5]
)

doc.add_paragraph()
add_para('Huong dan GitHub Pages:', bold=True, size=11, space_after=6)
gh_steps = [
    'Tao tai khoan tai github.com (mien phi)',
    'Tao repository moi, ten: kpi-hue-linh',
    'Upload 7 file: index.html, app.js, data_nhansu.js, manifest.json, sw.js, icon-192.png, icon-512.png',
    'Vao Settings -> Pages -> Source: main branch -> Save',
    'Cho 1-2 phut, app online tai: https://[username].github.io/kpi-hue-linh',
    'Chia se link nay cho nhan vien — ho cai PWA tren dien thoai nhu huong dan muc 3',
]
for i, s in enumerate(gh_steps, 1):
    add_para(f'  {i}. {s}', size=11, space_after=3)

doc.add_page_break()

# ========== 6. FAQ ==========
add_para('6. CAU HOI THUONG GAP (FAQ)', bold=True, size=14, color=(11,93,59), space_after=10)

faqs = [
    ('Du lieu co mat khi tat trinh duyet khong?', 'Khong. Du lieu luu trong localStorage, ton tai cho den khi xoa cache trinh duyet. Nhung nen SAO LUU dinh ky bang chuc nang Xuat JSON.'),
    ('Nhieu nguoi co the cung nhap tren 1 app khong?', 'Co, neu bat che do Google Sheets (TH2). Moi nguoi nhap tren dien thoai rieng, du lieu gui ve 1 Sheet chung.'),
    ('App co can Internet khong?', 'Khong (che do Cuc bo). Sau khi cai PWA, app chay offline 100%. Chi can Internet neu dung che do Google Sheets.'),
    ('Lam sao de khoi phuc du lieu khi doi dien thoai?', 'Xuat file JSON tren may cu (muc Sao luu), chuyen file sang may moi, vao app -> Nhap du lieu -> chon file JSON.'),
    ('Ai thay doi duoc danh muc may, nhan su?', 'Bat ky ai co quyen truy cap app deu co the sua trong muc "Nhan su & Danh muc". Nen chi cap cho quan doc/truong ca.'),
    ('Diem duoc tinh nhu the nao?', 'Diem nen A = Trung binh 5 tieu chi nhom A (thang 1-5). Da ky nang = Dem so chi so B+C dat 4 hoac 5 (toi da 6/6). Ket qua tong the tu dong de xuat dua tren diem A + da ky nang.'),
]
for q, a in faqs:
    add_para(f'H: {q}', bold=True, size=11, space_after=2)
    add_para(f'D: {a}', size=11, space_after=10)

# ========== 7. LUU Y ==========
add_para('7. LUU Y QUAN TRONG', bold=True, size=14, color=(11,93,59), space_after=10)
make_table(
    ['Van de', 'Giai phap'],
    [
        ['Mat du lieu khi xoa cache', 'Dinh ky Sao luu JSON (it nhat 1 lan/tuan)'],
        ['Khong thay 3 muc moi (BGC, KPI, Bao tri)', 'Xoa cache trinh duyet: Ctrl+Shift+R (PC) hoac xoa cache app (dien thoai)'],
        ['Muon nhieu nguoi nhap chung', 'Bat che do Google Sheets trong Cai dat'],
        ['Doi dien thoai / may tinh', 'Xuat JSON tu may cu -> Nhap vao may moi'],
        ['In phieu danh gia', 'Vao Lich su -> Xem phieu -> Nhan In phieu'],
        ['App khong cap nhat sau khi sua code', 'Xoa Service Worker: DevTools -> Application -> Service Workers -> Unregister'],
    ],
    col_widths=[5.5, 9]
)

doc.add_paragraph()
doc.add_paragraph()
add_para('--- Het ---', size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
add_para('Tai lieu nay duoc tao tu dong boi App Danh Gia Nang Luc CNSX — Phien ban V2.0', size=10, color=(100,100,100), align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('Cong ty TNHH TM-SX Hue Linh | Ngay: 17/06/2026', size=10, color=(100,100,100), align=WD_ALIGN_PARAGRAPH.CENTER)

# Save
output_path = r'D:\Cty Hue Linh\Danh Muc\Quy_Trinh_Danh_Gia_HL\App_Danh_Gia_KPI\Huong_Dan_Su_Dung_App_KPI_HueLinh.docx'
doc.save(output_path)
print(f'THANH CONG! File da luu tai:\n{output_path}')
